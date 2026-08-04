"""
Extracts plain text from uploaded study material (PDF, DOCX, TXT).
"""
import io

import docx
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class UnsupportedFileTypeError(Exception):
    pass


class TextExtractionError(Exception):
    pass


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_text = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages_text)
    except Exception as exc:  # pypdf raises various error types depending on the malformed input
        raise TextExtractionError(f"Could not read PDF: {exc}") from exc


def _extract_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        # Also pull text out of tables, which python-docx doesn't include in .paragraphs
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)
    except Exception as exc:
        raise TextExtractionError(f"Could not read DOCX: {exc}") from exc


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise TextExtractionError("Could not decode text file (unsupported encoding).")


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        text = _extract_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        text = _extract_docx(file_bytes)
    elif lower_name.endswith(".txt"):
        text = _extract_txt(file_bytes)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type for '{filename}'. Supported types: PDF, DOCX, TXT."
        )

    cleaned = _clean_text(text)
    if not cleaned.strip():
        raise TextExtractionError(
            "No extractable text was found in this file (it may be a scanned image without OCR)."
        )
    return cleaned


def _clean_text(text: str) -> str:
    # Collapse excessive blank lines/whitespace produced by PDF/DOCX extraction.
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)
